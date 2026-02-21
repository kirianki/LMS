
import docx
import json

def extract_docx_content(filepath):
    doc = docx.Document(filepath)
    content = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            content.append({
                'type': 'paragraph',
                'text': para.text,
                'style': para.style.name,
                'bold': any(run.bold for run in para.runs),
                'alignment': para.alignment
            })
            
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        content.append({
            'type': 'table',
            'data': table_data
        })
        
    return content

if __name__ == "__main__":
    filepath = "/home/sammy/Desktop/LMS/Mr.  GEORGE GITIBA NJENGA -OFFER LETTER BIASHARA INDIVIDUAL LOAN (2) (1).docx"
    try:
        content = extract_docx_content(filepath)
        print(json.dumps(content, indent=2))
    except Exception as e:
        print(f"Error: {e}")
