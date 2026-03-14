from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
from decimal import Decimal
from io import BytesIO
from apps.treasury.services.documents import get_tenant_branding

def format_currency(amount):
    """Format decimal amount as KES currency string."""
    try:
        val = float(amount)
        return f"KES {val:,.2f}"
    except (ValueError, TypeError):
        return "KES 0.00"

def create_document(tenant, title):
    """Create a branded document with header."""
    document = Document()
    branding = get_tenant_branding(tenant)
    
    # Title
    heading = document.add_heading(title.upper(), 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Institution Info
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(branding.get('company_name', 'Aurum Finance'))
    run.bold = True
    run.font.size = Pt(12)
    p.add_run(f"\n{branding.get('company_address', '')}")
    p.add_run(f"\n{branding.get('company_phone', '')} | {branding.get('company_email', '')}")
    
    document.add_paragraph().add_run().add_break() # Spacer
    
    return document

def add_footer(document, tenant):
    """Add simple footer."""
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = f"Generated for {get_tenant_branding(tenant).get('company_name', 'Aurum Finance')} on {datetime.date.today()}"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table(document, headers, data, bold_headers=True):
    """Add a styled table."""
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = str(header)
        if bold_headers:
             for paragraph in hdr_cells[i].paragraphs:
                 for run in paragraph.runs:
                     run.bold = True
    
    # Data
    for row_data in data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            
    return table

def generate_balance_sheet_docx(data, date, tenant):
    document = create_document(tenant, f"BALANCE SHEET (As of {date})")
    
    if not data.get('is_balanced', False):
        p = document.add_paragraph("WARNING: This Balance Sheet is Unbalanced!")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.bold = True

    def add_section(title, section_data):
        document.add_heading(title, level=2)
        table_data = []
        for item in section_data['details']:
            table_data.append([item['name'], format_currency(item['balance'])])
        
        # Total Row
        table_data.append(["", ""])
        table_data.append([f"Total {title}", format_currency(section_data['total'])])
        
        t = add_table(document, ["Account", "Balance"], table_data)
        
        # Bold last row
        rows = t.rows
        last_row = rows[-1]
        for cell in last_row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

    add_section("ASSETS", data['assets'])
    add_section("LIABILITIES", data['liabilities'])
    add_section("EQUITY", data['equity'])
    
    # Grand Total
    document.add_paragraph().add_run().add_break()
    total_le = Decimal(str(data['liabilities']['total'])) + Decimal(str(data['equity']['total']))
    
    summary_data = [
        ["Total Assets", format_currency(data['assets']['total'])],
        ["Total Liabilities & Equity", format_currency(total_le)]
    ]
    t = add_table(document, ["Metric", "Amount"], summary_data)
    
    add_footer(document, tenant)
    
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def generate_profit_loss_docx(data, start_date, end_date, tenant):
    document = create_document(tenant, "PROFIT & LOSS STATEMENT")
    document.add_paragraph(f"Period: {start_date} to {end_date}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_section(title, section_data):
        document.add_heading(title, level=2)
        if not section_data['details']:
            document.add_paragraph("No records found.")
            return

        table_data = []
        for item in section_data['details']:
            table_data.append([item['name'], format_currency(item['balance'])])
        
        table_data.append(["", ""])
        table_data.append([f"Total {title}", format_currency(section_data['total'])])
        
        t = add_table(document, ["Account", "Amount"], table_data)
        
        # Bold last row
        rows = t.rows
        last_row = rows[-1]
        for cell in last_row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

    add_section("INCOME", data['income'])
    add_section("EXPENSES", data['expenses'])
    
    document.add_paragraph().add_run().add_break()
    
    net_profit = Decimal(str(data['net_profit']))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"NET PROFIT / (LOSS): {format_currency(net_profit)}")
    run.bold = True
    run.font.size = Pt(14)
    if net_profit < 0:
        run.font.color.rgb = RGBColor(255, 0, 0)
    else:
        run.font.color.rgb = RGBColor(0, 128, 0)

    add_footer(document, tenant)
    
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def generate_cash_flow_docx(data, start_date, end_date, tenant):
    document = create_document(tenant, "CASH FLOW STATEMENT")
    p = document.add_paragraph(f"Period: {start_date} to {end_date}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n(Direct Method)")
    
    document.add_heading("OPERATING ACTIVITIES", level=2)
    
    # Inflows
    document.add_heading("Cash Inflows", level=3)
    if data['operating_activities']['inflow']:
        in_data = []
        for item in data['operating_activities']['inflow']:
            in_data.append([item.get('description', 'N/A'), item.get('date', ''), format_currency(item['amount'])])
        add_table(document, ["Description", "Date", "Amount"], in_data)
    else:
        document.add_paragraph("No inflows recorded.")
        
    document.add_paragraph()
    
    # Outflows
    document.add_heading("Cash Outflows", level=3)
    if data['operating_activities']['outflow']:
        out_data = []
        for item in data['operating_activities']['outflow']:
            out_data.append([item.get('description', 'N/A'), item.get('date', ''), f"({format_currency(item['amount'])})"])
        add_table(document, ["Description", "Date", "Amount"], out_data)
    else:
        document.add_paragraph("No outflows recorded.")
        
    document.add_paragraph()
    
    # Net Cash
    net_cash = Decimal(str(data['net_increase_in_cash']))
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"NET INCREASE / (DECREASE): {format_currency(net_cash)}")
    run.bold = True
    run.font.size = Pt(12)
    
    add_footer(document, tenant)
    
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def generate_generic_table_docx(title, data_list, headers, filters, tenant, summary=None):
    document = create_document(tenant, title)
    
    # Filters
    if filters:
        filter_text = []
        for k, v in filters.items():
            if v:
                filter_text.append(f"{k}: {v}")
        if filter_text:
            document.add_paragraph(" | ".join(filter_text)).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # Summary
    if summary:
        document.add_heading("Summary", level=2)
        sum_data = []
        for k, v in summary.items():
            val_str = str(v)
            if isinstance(v, (int, float, Decimal)) and 'count' not in k.lower():
               val_str = format_currency(v)
            sum_data.append([k.replace('_', ' ').title(), val_str])
        
        add_table(document, ["Metric", "Value"], sum_data)
        document.add_paragraph()
        
    # Main Table
    if data_list:
        table_headers = [h[1] for h in headers]
        table_data = []
        for item in data_list:
            row = []
            for h in headers:
                key = h[0]
                val = item.get(key, '')
                if 'amount' in key.lower() or 'balance' in key.lower() or 'principal' in key.lower() or 'interest' in key.lower():
                     try:
                         val = format_currency(val)
                     except:
                         pass
                row.append(str(val))
            table_data.append(row)
            
        add_table(document, table_headers, table_data)
    else:
        document.add_paragraph("No data found.")
        
    add_footer(document, tenant)
    
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def generate_trial_balance_docx(data, date, tenant):
    document = create_document(tenant, f"TRIAL BALANCE (As of {date})")
    
    table_data = []
    for item in data['report']:
        table_data.append([
            item['code'],
            item['name'],
            format_currency(item['debit']),
            format_currency(item['credit'])
        ])
    
    # Totals
    table_data.append([
        "", "TOTALS",
        format_currency(data['total_debit']),
        format_currency(data['total_credit'])
    ])
    
    t = add_table(document, ["Code", "Account", "Debit", "Credit"], table_data)
    
    # Bold last row
    rows = t.rows
    last_row = rows[-1]
    for cell in last_row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                
    add_footer(document, tenant)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def generate_general_ledger_docx(data, account_id, start_date, end_date, tenant):
    document = create_document(tenant, f"GENERAL LEDGER")
    document.add_paragraph(f"Period: {start_date} to {end_date}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_account_ledger(account_data):
        document.add_heading(f"Account: {account_data['account_code']} - {account_data['account_name']}", level=2)
        document.add_paragraph(f"Opening Balance: {format_currency(account_data['opening_balance'])}")
        
        table_data = []
        for item in account_data['history']:
            table_data.append([
                item['date'],
                item['description'],
                item['reference'],
                format_currency(item['debit']) if item['debit'] > 0 else "-",
                format_currency(item['credit']) if item['credit'] > 0 else "-",
                format_currency(item['balance'])
            ])
            
        add_table(document, ["Date", "Description", "Ref", "Debit", "Credit", "Balance"], table_data)
        
        document.add_paragraph()
        p = document.add_paragraph()
        run = p.add_run(f"Closing Balance: {format_currency(account_data['closing_balance'])}")
        run.bold = True
        document.add_paragraph() # Spacer

    if data.get('is_bulk'):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Consolidated Report Summary\nOpening: {format_currency(data['total_opening'])} | Closing: {format_currency(data['total_closing'])}")
        run.bold = True
        
        for acc in data['accounts']:
            add_account_ledger(acc)
    else:
        add_account_ledger(data)
    
    add_footer(document, tenant)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def generate_arrears_management_docx(data, tenant):
    """Generate Arrears Management Report DOCX."""
    document = create_document(tenant, "COLLECTIONS & ARREARS MANAGEMENT REPORT")
    document.add_paragraph(f"Report Date: {datetime.date.today().strftime('%B %d, %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. PAR Metrics Section
    document.add_heading("Portfolio at Risk (PAR) Metrics", level=2)
    par = data.get('par', {})
    par_data = [
        ["PAR 1+ Day", f"{par.get('par_1_plus_percent', 0):.2f}%", format_currency(par.get('par_1_plus_amount', 0))],
        ["PAR 30+ Days", f"{par.get('par_30_plus_percent', 0):.2f}%", format_currency(par.get('par_30_plus_amount', 0))],
        ["PAR 90+ Days", f"{par.get('par_90_plus_percent', 0):.2f}%", format_currency(par.get('par_90_plus_amount', 0))],
    ]
    add_table(document, ["Metric", "Percentage", "Outstanding Amount"], par_data)
    document.add_paragraph()

    # 2. Arrears Aging Buckets
    document.add_heading("Arrears Aging Buckets", level=2)
    aging = data.get('aging', {}).get('buckets', {})
    aging_data = []
    for bucket_name in ['current', '1-30', '31-60', '61-90', '90+']:
        b = aging.get(bucket_name, {'count': 0, 'amount': 0, 'balance': 0})
        aging_data.append([
            bucket_name.title(),
            str(b['count']),
            format_currency(b['amount']),
            format_currency(b['balance'])
        ])
    add_table(document, ["Bucket", "Loan Count", "Arrears Amount", "Portfolio Balance"], aging_data)
    document.add_paragraph()

    # 3. Active Cases Table
    document.add_heading("Detailed Collection Cases", level=2)
    cases = data.get('cases', [])
    if cases:
        case_headers = ["Loan #", "Borrower", "Days", "Amount", "Priority", "Status"]
        case_table_data = []
        for c in cases:
            case_table_data.append([
                str(c.get('loan_number', '')),
                str(c.get('borrower_name', 'N/A')),
                str(c.get('days_overdue', 0)),
                format_currency(c.get('overdue_amount', 0)),
                c.get('priority', '').upper(),
                c.get('status_display', '').title()
            ])
        add_table(document, case_headers, case_table_data)
    else:
        document.add_paragraph("No active collection cases found.")

    add_footer(document, tenant)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

